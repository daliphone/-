import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO
import os

# --- 1. 頁面配置與 UI ---
st.set_page_config(page_title="馬尼通訊 企劃提案系統 v14.3.8", page_icon="🐎", layout="centered")

st.markdown("""
    <style>
    .main { background-color: #F8FAFC; color: #1E293B; }
    [data-testid="stSidebar"] { background-color: #FFFFFF !important; border-right: 1px solid #E2E8F0 !important; }
    .section-header { 
        font-size: 20px !important; color: #003f7e !important; font-weight: 700 !important; 
        margin-top: 30px !important; margin-bottom: 10px !important;
        display: flex; align-items: center;
    }
    .section-header::before {
        content: ""; display: inline-block; width: 5px; height: 24px; 
        background-color: #ef8200; margin-right: 12px; border-radius: 2px;
    }
    .ai-btn-small>div>button { 
        background-color: #F5F3FF !important; color: #6D28D9 !important; 
        border: 1px solid #DDD6FE !important; font-size: 12px !important;
    }
    .version-info { font-size: 12px; color: #64748B; background: #F1F5F9; padding: 10px; border-radius: 8px; }
    </style>
    """, unsafe_allow_html=True)

# --- 2. 初始化 Session State (含可編輯的引導詞與建議) ---
FIELDS = ["p_name", "p_proposer", "p_purpose", "p_core", "p_schedule", "p_prizes", "p_sop", "p_marketing", "p_risk", "p_effect"]

# 預設的引導內容與建議內容
DEFAULT_LOGIC = {
    "p_purpose": "營運目的邏輯：強化解決痛點並增加商品銷售或去化高壓商品。",
    "p_core": "賣點配置建議：建立「低門檻、零風險」誘因。",
    "p_schedule": "執行重點建議：規劃宣傳、銷售、結案期資源分配。",
    "p_prizes": "配置用意：平衡大獎話題與小獎導流。",
    "p_sop": "執行注意事項：注入「卸下武裝」策略。",
    "p_marketing": "行銷策略：自動推薦管道並生成標語。",
    "p_risk": "風險管理：針對法務、稅務及損壞規範。",
    "p_effect": "成效效益：分析 O2O 轉換與名單累積。"
}

DEFAULT_TIPS = {
    "p_purpose": "核心：春節紅包議題，解決人流痛點。目標：引導消耗紅包財。",
    "p_core": "機制：購買禮包獲得序號。定價：$100 具備衝動購買力。",
    "p_schedule": "時程：1月中旬啟動，確保除夕前銷售完畢。",
    "p_prizes": "配置：PS5 (話題) + 現金。購物金用於官網引流。",
    "p_sop": "話術：先聊願望再推「試手氣」。SOP：強調序號正本。",
    "p_marketing": "宣傳：紅包視覺，社群任務設計分享好運。",
    "p_risk": "風險：每店配額管理。法規：中獎者身份證影本蒐集。",
    "p_effect": "指標：門市進店率、官網註冊數、轉化率。"
}

# 初始化 state
if 'logic_state' not in st.session_state: st.session_state.logic_state = DEFAULT_LOGIC.copy()
if 'tips_state' not in st.session_state: st.session_state.tips_state = DEFAULT_TIPS.copy()
if 'templates_store' not in st.session_state: st.session_state.templates_store = {"請選擇範本": {f: "" for f in FIELDS}}

for field in FIELDS:
    if field not in st.session_state: st.session_state[field] = ""

# --- 3. 頁面頂部：版本資訊 ---
with st.expander("ℹ️ 系統版本資訊 (v14.3.8)"):
    st.markdown("""
    <div class="version-info">
    <b>v14.3.8 (最新)</b>: 增加版本異動記錄；引導詞與實戰建議改為「可編輯模式」。<br>
    <b>v14.3.7</b>: 調整填寫版面為「直列順序」呈現。<br>
    <b>v14.3.6</b>: 導入清新視覺感、動態範本儲存與聯動系統。<br>
    <b>v14.3.5</b>: 模組化結構推進，整合「營運目的」與「去化高壓商品」邏輯。<br>
    <b>v14.3.4</b>: 建立基礎 AI 優化按鈕與 Word 導出功能。
    </div>
    """, unsafe_allow_html=True)

# --- 4. 側邊欄：範本與編輯模式切換 ---
with st.sidebar:
    st.header("📋 企劃管理")
    selected_tpl_key = st.selectbox("選擇既有範本", options=list(st.session_state.templates_store.keys()))
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📥 載入範本"):
            data = st.session_state.templates_store[selected_tpl_key]
            for k, v in data.items(): 
                if k in FIELDS: st.session_state[k] = v
            st.rerun()
    with col2:
        if st.button("💾 儲存範本"):
            if st.session_state.p_name:
                new_key = f"💾 {st.session_state.p_name[:10]}"
                st.session_state.templates_store[new_key] = {f: st.session_state[f] for f in FIELDS}
                st.success("儲存成功")
                st.rerun()
    
    st.divider()
    edit_mode = st.toggle("🔓 開啟引導詞/建議編輯模式", value=False)
    if edit_mode:
        st.info("編輯模式已開啟：您現在可以直接在下方的「邏輯提示」與「實戰建議」框內修改內容。")

# --- 5. 主要編輯區 (直列版面) ---
st.title("📱 模組化企劃系統 v14.3.8")

# 基本資訊
st.markdown('<p class="section-header">基本提案資訊</p>', unsafe_allow_html=True)
b1, b2, b3 = st.columns([2, 1, 1])
with b1: st.text_input("活動名稱", key="p_name")
with b2: st.text_input("提案人", key="p_proposer")
with b3: st.date_input("提案日期", value=datetime.now(), key="p_date")

st.divider()

# 章節配置定義
sections_info = [
    ("p_purpose", "一、 活動時機與目的"),
    ("p_core", "二、 活動核心內容"),
    ("p_schedule", "三、 活動時程安排"),
    ("p_prizes", "四、 贈品結構與預算"),
    ("p_sop", "五、 門市執行流程 (SOP)"),
    ("p_marketing", "六、 行銷流程與策略"),
    ("p_risk", "七、 風險管理與注意事項"),
    ("p_effect", "八、 預估成效")
]

# 直列渲染
for fid, title in sections_info:
    st.markdown(f'<p class="section-header">{title}</p>', unsafe_allow_html=True)
    
    # 1. 邏輯提示詞 (可編輯)
    if edit_mode:
        st.session_state.logic_state[fid] = st.text_input(f"修改「{title}」提示詞", value=st.session_state.logic_state[fid], key=f"edit_logic_{fid}")
    
    # 2. 填寫框 (帶入可編輯的提示詞作為 Placeholder)
    st.text_area("", key=fid, height=150, placeholder=st.session_state.logic_state[fid], label_visibility="collapsed")
    
    # 3. 功能區與實戰建議 (可編輯)
    c_ai, c_tip = st.columns([1, 4])
    with c_ai:
        if fid in ["p_purpose", "p_core", "p_marketing", "p_risk", "p_effect"]:
            st.markdown('<div class="ai-btn-small">', unsafe_allow_html=True)
            if st.button(f"🪄 AI 優化", key=f"btn_{fid}"):
                st.session_state[fid] = f"【AI 優化中】{st.session_state[fid]}"
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
    with c_tip:
        with st.expander("💡 查看/編輯實戰建議", expanded=False):
            if edit_mode:
                st.session_state.tips_state[fid] = st.text_area("編輯建議內容", value=st.session_state.tips_state[fid], key=f"edit_tip_{fid}", height=100)
            else:
                st.caption(st.session_state.tips_state[fid])
    st.write("") 

# --- 6. Word 產出 ---
def generate_word():
    doc = Document()
    doc.add_heading('行銷企劃執行提案書 v14.3.8', 0)
    doc.add_heading(st.session_state.p_name if st.session_state.p_name else "企劃書", level=1)
    for fid, title in sections_info:
        doc.add_heading(title, level=2)
        doc.add_paragraph(st.session_state[fid] if st.session_state[fid] else "（未填寫）")
    word_io = BytesIO(); doc.save(word_io); return word_io.getvalue()

st.divider()
if st.session_state.p_name:
    if st.button("✅ 完成企劃並產生文檔"):
        doc_data = generate_word()
        st.download_button(label="📥 下載企劃書", data=doc_data, file_name=f"MoneyMKT_{st.session_state.p_name}.docx")
